import streamlit as st
from redactor import redact_docx, redact_pdf
from io import BytesIO
from docx import Document
import re

st.set_page_config(page_title="Document Redaction AI", page_icon="")
st.title("AI Document Redaction Tool")
st.write("""
Upload a PDF or Word document to automatically redact sensitive information.  
You can also **manually unredact** any part and download the final document.
""")

uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])

def highlight_redacted(text):
    """Wrap [REDACTED] in HTML span to color it red"""
    highlighted_text = re.sub(r"\[REDACTED\]", r'<span style="color:red;font-weight:bold">[REDACTED]</span>', text)
    return highlighted_text

def save_text_to_docx(text):
    """Convert text into a Word document and return BytesIO"""
    out_doc = Document()
    for line in text.split("\n"):
        out_doc.add_paragraph(line)
    out_stream = BytesIO()
    out_doc.save(out_stream)
    out_stream.seek(0)
    return out_stream

if uploaded_file is not None:
    file_type = uploaded_file.type
    st.write(f"Processing file: {uploaded_file.name} ...")

    if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        redacted_file = redact_docx(uploaded_file)

        
        doc = Document(redacted_file)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        highlighted_text = highlight_redacted(full_text)

        st.subheader("Redacted Document Content (Highlighted):")
        st.markdown(highlighted_text.replace("\n", "  \n"), unsafe_allow_html=True)

        st.subheader("Edit / Unredact Text:")
        edited_text = st.text_area("You can replace [REDACTED] with original content here:", full_text, height=300)

        if st.button("Download Edited Document"):
            out_doc = save_text_to_docx(edited_text)
            st.download_button(
                label="Download Final Document",
                data=out_doc,
                file_name=f"final_{uploaded_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif file_type == "application/pdf":
        redacted_text = redact_pdf(uploaded_file)
        highlighted_text = highlight_redacted(redacted_text)

        st.subheader("Redacted PDF Content (Highlighted):")
        st.markdown(highlighted_text.replace("\n", "  \n"), unsafe_allow_html=True)

        st.subheader("Edit / Unredact Text:")
        edited_text = st.text_area("You can replace [REDACTED] with original content here:", redacted_text, height=300)

        if st.button("Download Edited Text as DOCX"):
            out_doc = save_text_to_docx(edited_text)
            st.download_button(
                label="Download Final Document",
                data=out_doc,
                file_name=f"final_{uploaded_file.name.replace('.pdf','.docx')}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:
        st.error("Unsupported file type.")
